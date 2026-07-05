import os
import time
import base64
import mimetypes
import logging
from typing import List, Tuple
import httpx
import pypdfium2 as pdfium
import pdfplumber
import fitz
import docx
import openpyxl
import xlrd
import csv
from io import BytesIO
from PIL import Image

logger = logging.getLogger("vigil.parsers")

# OpenRouter Free Vision Models Fallback Chain
OCR_MODEL_FALLBACKS = [
    "nvidia/nemotron-nano-12b-v2-vl:free",
    "google/gemma-4-26b-a4b-it:free",
    "openrouter/free"
]

def detect_document_type(file_path: str) -> Tuple[str, str]:
    """
    Detects document category (text-native vs scanned/image) and extension.
    Returns (category, extension).
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext in [".png", ".jpg", ".jpeg"]:
        return "image", ext
    elif ext == ".pdf":
        if is_pdf_scanned(file_path):
            return "scanned_pdf", ext
        else:
            return "text_native_pdf", ext
    elif ext in [".docx"]:
        return "text_native_docx", ext
    elif ext in [".xlsx", ".xls", ".csv"]:
        return "spreadsheet", ext
    else:
        # Fallback to general classification
        return "unknown", ext

def is_pdf_scanned(file_path: str) -> bool:
    """
    Checks if a PDF has sparse or zero extractable text (heuristic for scanned documents).
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for i, page in enumerate(pdf.pages):
                if i >= 3:  # Only sample first 3 pages
                    break
                page_text = page.extract_text()
                if page_text:
                    text += page_text.strip()
            return len(text) < 50
    except Exception as e:
        logger.warning(f"Error checking if PDF is scanned, defaulting to True: {str(e)}")
        return True

def parse_pdf_local(file_path: str) -> str:
    """
    Parses a text-native PDF locally using PyMuPDF (fitz) as primary,
    and falls back to pdfplumber if PyMuPDF fails or returns empty text.
    """
    # 1. Primary: PyMuPDF (fitz)
    try:
        logger.info(f"Attempting primary text extraction with PyMuPDF: {file_path}")
        text_pages = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_pages.append(page.get_text("text") or "")
        parsed_text = "\n\n--- Page Break ---\n\n".join(text_pages)
        
        # If PyMuPDF returned valid text (at least 10 non-whitespace chars), return it
        if parsed_text and len(parsed_text.strip()) >= 10:
            logger.info("Successfully extracted text using PyMuPDF.")
            return parsed_text
        else:
            logger.warning("PyMuPDF returned empty or too short text, falling back to pdfplumber.")
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {str(e)}. Falling back to pdfplumber.")

    # 2. Fallback: pdfplumber
    try:
        logger.info(f"Attempting fallback text extraction with pdfplumber: {file_path}")
        with pdfplumber.open(file_path) as pdf:
            text_pages = []
            for page in pdf.pages:
                text_pages.append(page.extract_text() or "")
            return "\n\n--- Page Break ---\n\n".join(text_pages)
    except Exception as e:
        raise Exception(f"pdfplumber extraction failed: {str(e)}")

def parse_docx_local(file_path: str) -> str:
    """
    Parses a DOCX document locally using python-docx.
    """
    try:
        doc = docx.Document(file_path)
        result = []
        for para in doc.paragraphs:
            if para.text.strip():
                result.append(para.text)
        
        for i, table in enumerate(doc.tables):
            result.append(f"\n### Table {i+1}")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                result.append("| " + " | ".join(row_text) + " |")
        
        return "\n\n".join(result)
    except Exception as e:
        raise Exception(f"docx extraction failed: {str(e)}")

def parse_xlsx_local(file_path: str) -> str:
    """
    Parses a modern Excel workbook locally using openpyxl.
    """
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        result = []
        for name in wb.sheetnames:
            sheet = wb[name]
            result.append(f"## Sheet: {name}\n")
            rows = list(sheet.iter_rows(values_only=True))
            rows = [r for r in rows if any(val is not None for val in r)]
            if not rows:
                continue
            
            headers = [str(h) if h is not None else "" for h in rows[0]]
            result.append("| " + " | ".join(headers) + " |")
            result.append("| " + " | ".join(["---"] * len(headers)) + " |")
            
            for row in rows[1:]:
                row_str = [str(cell) if cell is not None else "" for cell in row]
                if len(row_str) < len(headers):
                    row_str += [""] * (len(headers) - len(row_str))
                elif len(row_str) > len(headers):
                    row_str = row_str[:len(headers)]
                result.append("| " + " | ".join(row_str) + " |")
            result.append("")
        return "\n".join(result)
    except Exception as e:
        raise Exception(f"openpyxl extraction failed: {str(e)}")

def parse_xls_local(file_path: str) -> str:
    """
    Parses legacy Excel sheets locally using xlrd.
    """
    try:
        try:
            wb = xlrd.open_workbook(file_path)
        except Exception as e:
            if "xlsx file" in str(e).lower():
                import tempfile
                import shutil
                with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
                    tmp_name = tmp.name
                shutil.copyfile(file_path, tmp_name)
                try:
                    res = parse_xlsx_local(tmp_name)
                    return res
                finally:
                    if os.path.exists(tmp_name):
                        os.unlink(tmp_name)
            raise e
        result = []
        for name in wb.sheet_names():
            sheet = wb.sheet_by_name(name)
            result.append(f"## Sheet: {name}\n")
            if sheet.nrows == 0:
                continue
            
            rows = []
            for rx in range(sheet.nrows):
                row = [sheet.cell_value(rx, cx) for cx in range(sheet.ncols)]
                rows.append(row)
            
            rows = [r for r in rows if any(val != "" and val is not None for val in r)]
            if not rows:
                continue
            
            headers = [str(h) if h is not None else "" for h in rows[0]]
            result.append("| " + " | ".join(headers) + " |")
            result.append("| " + " | ".join(["---"] * len(headers)) + " |")
            
            for row in rows[1:]:
                row_str = [str(cell) if cell is not None else "" for cell in row]
                if len(row_str) < len(headers):
                    row_str += [""] * (len(headers) - len(row_str))
                elif len(row_str) > len(headers):
                    row_str = row_str[:len(headers)]
                result.append("| " + " | ".join(row_str) + " |")
            result.append("")
        return "\n".join(result)
    except Exception as e:
        raise Exception(f"xlrd extraction failed: {str(e)}")

def parse_csv_local(file_path: str) -> str:
    """
    Parses a CSV file locally using Python's csv module.
    """
    try:
        with open(file_path, mode='r', encoding='utf-8-sig', errors='ignore') as f:
            reader = csv.reader(f)
            rows = list(reader)
            if not rows:
                return ""
            
            result = []
            headers = rows[0]
            result.append("| " + " | ".join(headers) + " |")
            result.append("| " + " | ".join(["---"] * len(headers)) + " |")
            
            for row in rows[1:]:
                result.append("| " + " | ".join(row) + " |")
            return "\n".join(result)
    except Exception as e:
        raise Exception(f"CSV extraction failed: {str(e)}")

def parse_unstructured_local(file_path: str) -> str:
    """
    Fallback parser using Unstructured.
    """
    try:
        from unstructured.partition.auto import partition
        elements = partition(filename=file_path)
        return "\n\n".join([str(el) for el in elements])
    except Exception as e:
        raise Exception(f"Unstructured extraction fallback failed: {str(e)}")

def parse_via_openrouter_ocr(file_path: str, api_key: str) -> Tuple[str, str]:
    """
    Parses scanned image files or PDFs using the OpenRouter Vision API.
    Attempts model fallback sequence.
    Returns (transcription, model_used).
    """
    if not api_key:
        raise Exception("OpenRouter API key is missing. Set OPENROUTER_API_KEY in .env.")
        
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    # 1. Convert Scanned PDFs to images
    base64_images: List[str] = []
    if ext == ".pdf":
        doc = pdfium.PdfDocument(file_path)
        for i, page in enumerate(doc):
            bitmap = page.render(scale=2)  # Scale=2 for clear OCR
            pil_img = bitmap.to_pil()
            buffered = BytesIO()
            pil_img.save(buffered, format="PNG")
            img_b64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
            base64_images.append(img_b64)
    else:
        with open(file_path, "rb") as image_file:
            img_b64 = base64.b64encode(image_file.read()).decode("utf-8")
            base64_images.append(img_b64)
            
    # 2. Call OpenRouter with fallback models
    last_error = None
    for model in OCR_MODEL_FALLBACKS:
        try:
            logger.info(f"Attempting OCR using model: {model}")
            transcriptions = []
            
            for idx, base64_image in enumerate(base64_images):
                if len(base64_images) > 1:
                    logger.info(f"Transcribing page {idx+1}/{len(base64_images)}")
                
                # Execute API request with backoff
                text = run_api_ocr_request(base64_image, model, api_key)
                transcriptions.append(text)
                
            return "\n\n--- Page Break ---\n\n".join(transcriptions), model
            
        except Exception as e:
            logger.warning(f"Model {model} failed: {str(e)}. Attempting next fallback model...")
            last_error = e
            continue
            
    raise Exception(f"All OCR fallback models failed. Last error: {str(last_error)}")

def run_api_ocr_request(base64_image: str, model_slug: str, api_key: str) -> str:
    """
    Dispatches the base64 encoded image to OpenRouter with rate-limit exponential backoff.
    """
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model_slug,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Transcribe all text from this image exactly. Keep formatting, tables, and layouts intact where possible. Output direct transcripts only."
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{base64_image}"
                        }
                    }
                ]
            }
        ]
    }
    
    # 3 attempts with backoff
    for attempt in range(3):
        try:
            with httpx.Client(timeout=90.0) as client:
                response = client.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers=headers,
                    json=payload
                )
                if response.status_code == 200:
                    data = response.json()
                    return data["choices"][0]["message"]["content"]
                elif response.status_code == 429:
                    wait_time = (2 ** attempt) * 2
                    logger.warning(f"Rate limited (429). Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    raise Exception(f"HTTP {response.status_code}: {response.text}")
        except Exception as e:
            if attempt == 2:
                raise e
            wait_time = (2 ** attempt) * 2
            logger.warning(f"Request failed: {str(e)}. Retrying in {wait_time}s...")
            time.sleep(wait_time)
            
    raise Exception("Max OCR API retries exceeded.")
